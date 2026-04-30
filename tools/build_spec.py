# -*- coding: utf-8 -*-
"""Build JLPT N5 Grammar Tutor Functional Specification as a .docx file.

Run from the repo root:
    python tools/build_spec.py

Overwrites: JLPT N5 Grammar Tutor - Functional Spec.docx
Requires:   pip install python-docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "specifications"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = str(OUT_DIR / "JLPT N5 Grammar Tutor – Functional Spec.docx")

doc = Document()

# ---------- Style configuration ----------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Tighter heading sizes
for lvl, size in [(1, 16), (2, 13), (3, 11)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor(0x14, 0x45, 0x2A)


# ---------- Helpers ----------
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para

def bullet(text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    runs = []
    if '**' in text:
        # split on bold markers
        parts = text.split('**')
        for i, part in enumerate(parts):
            r = para.add_run(part)
            r.bold = (i % 2 == 1)
            runs.append(r)
    else:
        para.add_run(text)
    return para

def numbered(text):
    para = doc.add_paragraph(style='List Number')
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            r = para.add_run(part)
            r.bold = (i % 2 == 1)
    else:
        para.add_run(text)
    return para

def code_block(text):
    for line in text.splitlines():
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line if line else " ")
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    # spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)

def hr():
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_run = title_p.add_run("Functional Specification - JLPT N5 Grammar Tutor & Test")
title_run.font.size = Pt(20)
title_run.bold = True
title_run.font.color.rgb = RGBColor(0x14, 0x45, 0x2A)

sub_p = doc.add_paragraph()
sub_run = sub_p.add_run("(GitHub-Hosted Static Web App)  •  Version 4  •  Amended 2026-04-30")
sub_run.italic = True
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p("Target learner: late-N5 (wrapping up N5; preparing for JLPT N5 exam).")
p("Live site: https://gauravaccentureproducts.github.io/jlpt-n5-tutor/")
p("Predecessors: v1 (KnowledgeBank/sources.md era) and v2 (post-Brief-1) preserved as separate files.")

hr()

# ---------- Changelog from v2 ----------
h1("0) Changelog from v2")
p("Shipped between 2026-04-29 and 2026-04-30 across UX Brief 2 (4 phases), the Pass-9 content correction, the Pass-10 audio fixes, and a battery of UI follow-ups.")

h2("New routes & flow (Brief 2 §1, §14)")
bullet("**`#/home`** is now the default route. First-time visitors see a CTA + 3 pillar cards; returning visitors see a Continue card + Today's review queue + 7-day streak heatmap.")
bullet("**`#/learn`** is now a 5-card hub (Reference: Grammar / Vocabulary / Kanji ; Practice: Dokkai / Listening). 3+2 semantic split avoids orphaning.")
bullet("**`#/learn/grammar`** - 187 grammar cards in 32 categories with a sticky chip-rail jump menu.")
bullet("**`#/learn/vocab`** - 1002 vocab cards in 40 sections (only first section open by default) with a 40-chip jump menu that opens + scrolls the target section.")
bullet("**`#/learn/vocab/<form>`** - per-word detail page (form / reading / gloss / 5 example sentences sourced from grammar.json by either kanji form OR kana reading).")
bullet("**`#/kanji`** - card grid (106 entries with glyph + meaning + first kun/on); **`#/kanji/<glyph>`** - per-kanji detail page with prev/next nav.")
bullet("**`#/test/<n>`** for n in {20,30,50} starts a test of that length directly.")

h2("Settings panel additions (Brief 2 §5)")
bullet("**3-mode furigana** (Always show / Hide on known kanji / Never) replaces the binary toggle. Live preview next to the radios.")
bullet("**Per-kanji 'I know this' flag** - click any kanji glyph anywhere to open a popover with on/kun/meaning + a checkbox; flag persists in `localStorage.knownKanji` and feeds the 'hide-known' furigana mode.")
bullet("**Audio playback speed** (0.75x / 1.0x / 1.25x) - applied via MutationObserver to every `<audio>` element on every route render.")
bullet("**Reduce motion** override (auto / on / off) - sets `data-reduce-motion` on `<html>`; CSS overrides animation/transition durations.")
bullet("**Typed-phrase reset** ('Type RESET to confirm') replaces the legacy double-confirm dialog.")
bullet("**Export schema bumped 1 → 2** to include the new `knownKanji` and `streak` keys (settings, history, results, knownKanji, streak). v1 imports remain accepted - missing keys default to empty.")

h2("PWA / offline upgrades (Brief 2 §12)")
bullet("**Service worker `jlpt-n5-tutor-v24`** - stale-while-revalidate for the shell (HTML/CSS/JS) and cache-first for content (data/audio/locales).")
bullet("**'Update available' toast** - SW posts a SW_UPDATE_AVAILABLE message when new shell bytes are detected; click Reload to skipWaiting + reload.")
bullet("**PWA install banner** - one-time, dismissible, persists in localStorage.")
bullet("**Offline indicator** - chip in top-right that toggles with `navigator.onLine`. Hidden when online.")

h2("Mobile / responsive / a11y (Brief 2 §9, §10)")
bullet("**Bottom nav at ≤480 px** - primary nav becomes a fixed bar with `env(safe-area-inset-bottom)` honored.")
bullet("**44 px tap-target floor** enforced via CSS on every interactive element.")
bullet("**Skeleton screens** replace the legacy 'Loading...' text on every route. 5-second timeout shows a real 'Couldn't load - Retry' UI.")
bullet("**Empty states with routing buttons** for Review / Summary / Test (no completed tests / no progress yet / no due items).")
bullet("**Print stylesheet** (`@media print`) hides chrome, expands `<details>`, switches to serif. Clean Learn-lesson printables.")

h2("Power user / search / shortcuts (Brief 2 §7, §8)")
bullet("**Cross-corpus search** in the secondary nav: indexes grammar (id/pattern/meaning/explanation), vocab (form/reading/gloss, capped at 1500), kanji (glyph/on/kun/meanings). Results group by type with counts. `/` keyboard shortcut focuses the input.")
bullet("**Global keyboard shortcuts** (`js/shortcuts.js`): 1-4 picks Nth choice; Space reveal/flip; Enter clicks primary/Submit; ? opens cheatsheet overlay; Esc dismisses overlays.")
bullet("**Quit-prompt in Test** - `__testInProgress` flag set on test start; beforeunload + hashchange interceptors confirm 'Quit this test?' and revert the hash on cancel.")

h2("Audio + content corrections (Pass 9, Pass 10)")
bullet("**491 MP3 files** - 449 grammar examples + 30 reading passages + 12 listening scripts (~19 MB) generated via gTTS at build time. App degrades gracefully when MP3s are absent.")
bullet("**`tools/build_audio.py:normalize_for_tts()`** - converts ASCII digits to kanji digits before passing to gTTS so '3さつ' reads as 'さんさつ' and not as 'スリーさつ'. 274 occurrences across the corpus normalized.")
bullet("**`data/n5_kanji_readings.json`** - 35 wrong-primary readings fixed (本 もと→ほん, 時 とき→じ, 月 つき→がつ, 学 まな→がく, 人 ひと→にん, etc.) so auto-furigana renders the N5-context-most-common reading.")
bullet("**Pass 9 (external content brief)** - 27 + 4 sweeps + 7 cross-file consistency checks raised, all 38 fixed, 0 open across moji/goi/bunpou/dokkai/authentic question files.")
bullet("**Pass 10 (audio + auto-furigana)** - 309 findings (274 ASCII-digit + 35 wrong-primary), all closed.")

h2("Cumulative content audit tally")
bullet("**Pass 1-9** content corrections: 153 raised, 153 closed.")
bullet("**Pass 10** audio + auto-furigana: 309 raised, 309 closed.")
bullet("**Total: 462 raised, 462 closed, 0 open** across 10 audit passes spanning JLPT paper-maker, native-teacher, external 日本語教師, and TTS / rendering correctness perspectives.")

h2("Repo / specifications hygiene")
bullet("**`specifications/`** directory holds the canonical functional spec (this document).")
bullet("**`not-required/`** directory archives the v1 spec and other transient artifacts (kept on disk but out of the working tree).")
bullet("**`CHANGELOG.md`** documents user-visible changes per release; linked from the footer's 'What's new' link.")
bullet("**Em-dash-free codebase** - 881 occurrences stripped from shipped files for cp932 console safety. Regression guard X-6.5 in the testing plan.")

hr()

# ---------- Changelog from v1 ----------
h1("0.1) Changelog from v1 (historical)")
bullet("**§4.2 Furigana default flipped to OFF on N5 kanji** - matches a late-N5 learner. Per-session toggle to flip ON when desired.")
bullet("**§4.3 Lesson template expanded from 5 blocks to 7** - adds Form & Connection Rules and Common Mistakes / Contrasts blocks.")
bullet("**§4.4 NEW Canonical pattern source** - locked to local file `KnowledgeBank/grammar_n5.md` (23 categories, exhaustive), not an external website.")
bullet("**§4.5 NEW meaning_ja rule** - N5 kanji allowed; no out-of-scope kanji or vocab; no romaji anywhere.")
bullet("**§4.6 NEW Example variety rule** - 4-6 examples for fundamentals with mandatory form variety (`affirmative` / `negative` / `question` / `past` / `dialogue`).")
bullet("**§4.7 NEW N5 vocabulary whitelist** - content lint enforces scope using `KnowledgeBank/vocabulary_n5.md`.")
bullet("**§4.8 NEW N5 kanji whitelist** - furigana rule and content lint enforce scope using `KnowledgeBank/kanji_n5.md`.")
bullet("**§5.3 Submit button rule clarified** - visible-but-disabled with tooltip and remaining-count, not hidden.")
bullet("**§5.7 NEW Diagnostic test** - optional 10-question placement check on first run.")
bullet("**§5.8 NEW Drill mode (SRS-light)** - promoted from Phase 2 roadmap to Phase 1. Intervals 1d / 3d / 7d / 14d.")
bullet("**§6.2 Test engine** - adds `sentence_order` (並べ替え) question type.")
bullet("**§6.5 NEW Per-distractor explanations** for high-confusion questions.")
bullet("**§6.6 NEW Threshold-based weak pattern detection** - `errors/attempts ≥ 0.5 AND attempts ≥ 2` over rolling history, not single-test snapshot.")
bullet("**§7.1 / §7.2 schemas extended** - pattern adds `category`, `order`, `form_rules`, `common_mistakes`; question adds `direction`, `distractor_explanations`, `high_confusion`, `sentence_order` shape.")
bullet("**§8 Tech stack locked** - Vanilla HTML + CSS + JavaScript with ES modules. No build step. No npm. Open `index.html` to run.")
bullet("**§9 Repo structure updated** - co-locates app with existing content files in `/JLPT/N5/`.")

hr()

# ---------- 1) Purpose & Goals ----------
h1("1) Purpose & Goals")
p("Build a PC-friendly, browser-based web app that:")
bullet("Teaches all JLPT N5 grammar patterns in a structured Chapter 1 format (TOC → meaning → form/connection → explanation → examples → common mistakes → simple Japanese meaning).")
bullet("Provides an auto-graded test (Chapter 2) attemptable on a PC.")
bullet("Enforces: Submit button visible-but-disabled until all questions are answered; results display immediately after Submit (no reload).")
bullet("Performs gap analysis using rolling-history thresholds (Chapter 3) and reteaches weak patterns with form rules, common mistakes, and fresh examples.")
bullet("Drives a Drill mode (Spaced Repetition) over weak patterns at 1d / 3d / 7d / 14d intervals.")
bullet("Summarizes mastered / weak / untested patterns and recommends next steps (Chapter 4).")
bullet("Runs on GitHub Pages as a static site with zero build step. Anyone with a modern browser can clone and double-click `index.html`.")

# ---------- 2) Personas ----------
h1("2) Target Users & Personas")
bullet("**Learner (Krivi):** Wrapping up N5; preparing for the JLPT N5 exam. Already knows most N5 kanji visually. Default mode hides furigana on N5 kanji.")
bullet("**Tutor / Admin (You):** Maintains grammar content + question bank in source-of-truth files (`KnowledgeBank/grammar_n5.md`, `KnowledgeBank/vocabulary_n5.md`, `KnowledgeBank/kanji_n5.md`, `/data/grammar.json`, `/data/questions.json`). Reviews progress.")

# ---------- 3) Scope ----------
h1("3) Scope")
h2("In Scope (Phase 1)")
bullet("Grammar lessons (TOC + 7-block pattern detail pages).")
bullet("Auto-graded test engine (MCQ + dropdown fill + sentence ordering 並べ替え).")
bullet("Instant scoring with per-question review and per-distractor explanations.")
bullet("Auto-generated weak grammar pattern list using rolling threshold detection.")
bullet("Drill mode (SRS-light) over weak patterns.")
bullet("Diagnostic placement test on first run.")
bullet("Local progress tracking via LocalStorage (required, not optional).")
bullet("GitHub Pages deployment with no build step.")
bullet("Furigana toggle (per-session) for N5-kanji visibility.")
bullet("Content lint script (run by author): flags non-N5 vocab/kanji in `/data/*.json`.")

h2("Out of Scope (Phase 1)")
bullet("User accounts / login.")
bullet("Server-side analytics.")
bullet("Open-ended writing or essay evaluation.")
bullet("Audio / listening practice.")
bullet("Adaptive testing with Anki-style ease factors (Drill is Leitner-light only).")
bullet("Multi-learner profiles.")

# ---------- 4) Content Rules ----------
h1("4) Content Rules (Hard Requirements)")

h2("4.1 JLPT N5 Scope Constraints")
bullet("All lesson examples and test content must use only N5-scope vocabulary (per `KnowledgeBank/vocabulary_n5.md`) and N5-scope kanji (per `KnowledgeBank/kanji_n5.md`).")
bullet("If out-of-scope text is unavoidable: prefer N5 alternative; otherwise apply furigana rule (§4.2) to the specific word only.")
bullet("**No romaji anywhere** in lesson or test content. Force kana reading.")

h2("4.2 Furigana Rules (Strict)")
bullet("**Default: furigana OFF on N5-scope kanji.** This matches a late-N5 learner; the goal is reading practice, not crutches.")
bullet("**Always furigana ON** for any kanji or word beyond N5 scope, on the specific word only - never the entire sentence.")
bullet("**Per-session toggle:** \"Show furigana on N5 kanji\" - learner can flip ON for extra support; OFF is the default and persists in LocalStorage.")
bullet("Furigana rendering must be a true ruby annotation (`<ruby><rt>`), not parenthetical. Must support per-character or per-word annotation, not whole-sentence.")

h2("4.3 Lesson Template (Chapter 1) - 7 Blocks")
p("For each grammar pattern, display in this exact order:")
numbered("**Pattern name** - e.g., 〜です／〜ます")
numbered("**Meaning (English)** - short gloss.")
numbered("**Form & connection rules** - what does the pattern attach to (noun / い-adj / な-adj / verb-which-form)? Include a small conjugation table where applicable (affirmative / negative / past / past-negative / question).")
numbered("**Explanation & usage (English)** - when to use it, register, nuance.")
numbered("**Example sentences (4-6 for fundamentals; 2-3 for narrow patterns)** - must follow the form-variety rule (§4.6).")
numbered("**Common mistakes / contrasts** - pairs like 「は vs が」, 「に vs で」, 「あります vs います」, 「上手 vs 得意」. At least one entry where applicable. Include a strikethrough wrong sentence with one-line \"why this is wrong\" where helpful.")
numbered("**Meaning (Japanese, simple)** - N5 kanji allowed, no out-of-scope words, hiragana otherwise. Provides reading reinforcement.")

h2("4.4 Canonical Pattern Source")
bullet("The canonical pattern coverage list is the local file **`KnowledgeBank/grammar_n5.md`** (in this directory) - 23 categories, exhaustive at N5 level.")
bullet("TOC ordering follows that file's 23-category structure: copula → particles → demonstratives → question words → verbs (ます-form, plain form) → te-form → adjectives → existence → comparison → desiderative → counters → time → conjunctions → giving/receiving → causation → nominalization → set patterns → frequency → polite phrases → other core patterns → honorifics.")
bullet("Every pattern listed in `KnowledgeBank/grammar_n5.md` MUST have a matching entry in `/data/grammar.json`. A coverage-check script (`tools/check_coverage.py`) enforces this - fails commit if a pattern is missing.")
bullet("Authoritative reference sources for content quality are documented in **`KnowledgeBank/sources.md`** (Genki, Minna no Nihongo, Try!, Makino & Tsutsui, Marugoto, Bunpro, etc.).")

h2("4.5 meaning_ja Rule")
bullet("N5 kanji ALLOWED for content nouns/verbs.")
bullet("All other words MUST be hiragana.")
bullet("No out-of-scope kanji or vocabulary.")
bullet("No romaji.")
bullet("If an out-of-scope word is unavoidable, apply furigana to that word only.")

h2("4.6 Example Variety Rule")
p("For fundamental patterns (high frequency: は、が、を、に、〜ます、〜ました、〜てください、〜たい、etc.), provide 4-6 examples with form variety:")
bullet("At least 1 affirmative.")
bullet("At least 1 negative.")
bullet("At least 1 question.")
bullet("At least 1 past tense (where applicable).")
bullet("Optionally 1 mini-dialogue.")
p("Each example MUST declare a `form` field: `affirmative` | `negative` | `question` | `past` | `dialogue`.")
p("For narrow patterns (e.g., specific particles like 「や」), 2-3 examples are acceptable.")

h2("4.7 N5 Vocabulary Whitelist")
bullet("`KnowledgeBank/vocabulary_n5.md` is the canonical N5 vocabulary scope list.")
bullet("Generated artifact `/data/n5_vocab_whitelist.json` is consumed by the content lint script.")
bullet("Lint script flags any token in `/data/grammar.json` and `/data/questions.json` that is not on the whitelist (excluding hiragana-only function words / particles / explicit furigana-tagged out-of-scope words).")

h2("4.8 N5 Kanji Whitelist")
bullet("`KnowledgeBank/kanji_n5.md` is the canonical N5 kanji scope list.")
bullet("Generated artifact `/data/n5_kanji_whitelist.json` is consumed by the furigana renderer and the lint script.")
bullet("The furigana renderer treats any kanji NOT on this whitelist as out-of-scope and renders furigana for it; kanji ON this whitelist render furigana only when the user toggle is ON.")

# ---------- 5) UX / UI ----------
h1("5) UX / UI Requirements")

h2("5.1 Global Navigation")
p("Primary navigation tabs (in this order):")
bullet("**Learn** (Chapter 1)")
bullet("**Test** (Chapter 2)")
bullet("**Drill** (NEW - SRS-driven mini-sessions on weak patterns; shows count badge of items due today)")
bullet("**Review Weak Areas** (Chapter 3)")
bullet("**Summary** (Chapter 4)")

h2("5.2 Chapter 1 - Learn")
bullet("TOC page lists all N5 grammar patterns grouped by category (per `KnowledgeBank/grammar_n5.md`'s 23 categories).")
bullet("Each TOC entry shows pattern name + a small \"known / weak / untested\" badge derived from rolling history.")
bullet("Clicking a pattern opens the 7-block detail view (§4.3).")
bullet("Provide Next / Previous pattern navigation within and across categories.")
bullet("Provide search/filter by keyword (pattern name, meaning, or category).")
bullet("Provide a \"Mark as known\" checkbox per pattern (manual override; updates rolling history).")

h2("5.3 Chapter 2 - Test")
h3("Test Setup")
bullet("Default test length: **20** questions. Selectable: 20 / 30 / 50.")
bullet("First-time learner: a one-time prompt offers the Diagnostic Test (§5.7) before the first real test.")
bullet("Question types (auto-gradable):")
bullet("  Multiple choice (single correct).", level=1)
bullet("  Fill-in-blank with dropdown (preferred for particles).", level=1)
bullet("  Sentence ordering (並べ替え) - 4-5 word/phrase tiles, click to order.", level=1)
bullet("  Optional: strict short input (exact match with normalization rules).", level=1)
bullet("Question sampling MUST balance across categories (no single-category dominance), unless test length < 8.")

h3("Attempting")
bullet("Progress indicator: \"Question X of N\" + a thin progress bar.")
bullet("Allow Next / Previous navigation.")
bullet("Optional \"Flag question\" feature (icon toggle; flagged items grouped at end of review).")

h3("Submit Button Rule (Mandatory)")
bullet("Submit button is **always visible** on the last question.")
bullet("Until every question has an answer, Submit is **disabled** and shows a tooltip on hover: \"Answer all questions to submit\" plus a remaining-count: \"3 questions still unanswered\".")
bullet("Once all questions are answered, Submit becomes enabled (no extra confirmation step).")
bullet("This rule is non-negotiable across all question types and test lengths.")

h3("Post-Submit Behavior (Mandatory)")
bullet("Clicking Submit triggers immediate in-place results rendering.")
bullet("No manual grading.")
bullet("No page reload.")
bullet("Test responses are appended to rolling history in LocalStorage before the results screen renders.")

h2("5.4 Results Screen (After Submit)")
p("Three sections, in this order:")

h3("(a) Score Summary")
bullet("Total score, e.g., 17/20.")
bullet("Correct count / incorrect count.")
bullet("Percentage.")
bullet("Time taken (optional).")

h3("(b) Answer Review (per question)")
bullet("Question text.")
bullet("Learner's answer (highlighted; red if incorrect, green if correct).")
bullet("Correct answer (clearly shown).")
bullet("Explanation of why the correct answer is correct.")
bullet("**Per-distractor explanation** for high-confusion questions: only the distractor the learner picked is explained (not all four), to avoid cognitive overload.")
bullet("Tagged grammar pattern(s) with link to Chapter 1 detail page.")

h3("(c) Grammar Gap List")
bullet("Patterns flagged as weak by rolling-history threshold (§6.6): `errors/attempts ≥ 0.5 AND attempts ≥ 2`.")
bullet("Each item links to Chapter 3 (Review Weak Areas).")
bullet("Each weak pattern is auto-queued into Drill at the 1-day interval.")

h2("5.5 Chapter 3 - Review Weak Areas")
p("For each weak grammar pattern, display:")
bullet("Pattern name.")
bullet("Meaning + explanation (English).")
bullet("**Form & connection rules** + **common mistakes** (re-shown with extra emphasis).")
bullet("Why the wrong answer was wrong (per-distractor explanations pulled from missed questions).")
bullet("New examples (4-6 if fundamental, 2-3 otherwise) - must be DIFFERENT from Chapter 1 examples to avoid memorization-by-position.")
bullet("\"Drill these now\" button → opens an SRS Drill session focused on this pattern.")

h2("5.6 Chapter 4 - Summary")
bullet("**Mastered patterns** - seen ≥ 2 times AND error rate < 30% over rolling history (or manually marked known).")
bullet("**Patterns needing more practice** - `errors/attempts ≥ 0.5 AND attempts ≥ 2`.")
bullet("**Untested patterns** - seen 0-1 times → \"Take a test to assess.\"")
bullet("**Suggested next steps:** Drill due items / Re-attempt test on weak patterns / Take Diagnostic if not yet done.")
bullet("Visual: a heatmap-style grid by category showing strength (green) / weak (red) / untested (gray).")

h2("5.7 NEW: Diagnostic Test (First Run)")
bullet("Prompt on first visit (and re-runnable from Settings): \"Take a 10-question diagnostic to map your current strengths?\"")
bullet("10 questions, sampled across the 23 categories of `KnowledgeBank/grammar_n5.md` with priority on the highest-frequency categories (copula, particles, verbs, adjectives, te-form, tai-form, counters, conjunctions, demonstratives, question words).")
bullet("Results don't count toward score history but DO seed the initial weak-pattern list and Drill queue.")
bullet("Skippable - learner can go straight to lessons or full test.")

h2("5.8 NEW: Drill Mode (SRS-light)")
bullet("Driven by LocalStorage: per-pattern SRS state.")
bullet("Intervals: **1 day / 3 days / 7 days / 14 days** (Leitner-style boxes).")
bullet("A pattern enters Drill on:")
bullet("  First miss in a Test or Diagnostic.", level=1)
bullet("  Any wrong answer in a previous Drill session.", level=1)
bullet("A pattern advances to the next interval on a correct answer; resets to 1d on wrong.")
bullet("A pattern graduates (leaves Drill) after 4 consecutive correct answers across boxes.")
bullet("UI: \"Drill\" tab shows a count badge of patterns due today.")
bullet("A Drill session = 5-10 questions sampled from due patterns. Feedback is **immediate per question** (not batched like the full Test).")
bullet("Drill respects the visible-but-disabled Submit rule for any multi-question Drill batch ≥ 5 questions.")

# ---------- 6) Functional Requirements ----------
h1("6) Functional Requirements (Detailed)")

h2("6.1 Grammar Content Management")
bullet("**FR-L1:** System shall load `/data/grammar.json` at startup.")
bullet("**FR-L2:** System shall render TOC grouped by category (per `KnowledgeBank/grammar_n5.md`'s 23 categories), in declared order.")
bullet("**FR-L3:** System shall render pattern detail using the 7-block expanded template (§4.3).")
bullet("**FR-L4:** System shall support cross-linking from Test results → Chapter 1 pattern detail and Chapter 3 review.")
bullet("**FR-L5:** System shall enforce that every pattern in `KnowledgeBank/grammar_n5.md` has a matching entry in `/data/grammar.json`. The `tools/check_coverage.py` script verifies this and exits non-zero if any pattern is missing.")

h2("6.2 Test Engine")
bullet("**FR-T1:** Generate tests from `/data/questions.json`. Each question requires: `id`, `grammarPatternId`, `type`, `correctAnswer`, `choices` (for MCQ/dropdown) OR `tiles`+`correctOrder` (for sentence_order). Recommended: `explanation_en`, `distractor_explanations`.")
bullet("**FR-T2:** Enforce \"all answered before Submit becomes enabled\" via visible-but-disabled rule (§5.3).")
bullet("**FR-T3:** Compute score instantly on Submit.")
bullet("**FR-T4:** Display results immediately after Submit (no reload).")
bullet("**FR-T5:** Produce gap list using rolling threshold (§6.6), not single-test snapshot.")
bullet("**FR-T6:** Support `sentence_order` question type with array-equality grading.")
bullet("**FR-T7:** Sample test questions to balance across categories - no single category may exceed `ceil(N / 5)` questions in a test of length N (where N ≥ 8).")

h2("6.3 Answer Normalization (if short input is used)")
bullet("**FR-N1:** Trim whitespace; convert full-width ↔ half-width as appropriate.")
bullet("**FR-N2:** Support multiple acceptable answers (array) where defined.")
bullet("**FR-N3:** Normalize hiragana / katakana variants where the question allows either.")

h2("6.4 Progress & Persistence (REQUIRED)")
bullet("**FR-P1:** Persist test results, weak patterns, SRS state, and rolling pattern history in LocalStorage.")
bullet("**FR-P2:** Provide a \"Reset progress\" button with a two-step confirmation.")
bullet("**FR-P3:** Persist learner's furigana toggle preference and last-selected test length.")
bullet("**FR-P4:** Use the LocalStorage key namespace `jlpt-n5-tutor:*` to avoid collisions.")

h2("6.5 NEW: Per-Distractor Explanations")
bullet("**FR-D1:** For questions with `high_confusion: true`, every distractor MUST have an explanation in `distractor_explanations`.")
bullet("**FR-D2:** Results screen shows the distractor explanation only for the choice the learner picked, not all distractors - to avoid information overload.")
bullet("**FR-D3:** Chapter 3 (Review Weak Areas) aggregates distractor explanations from all missed questions for the same pattern.")

h2("6.6 NEW: Weak Pattern Detection (Threshold-Based)")
bullet("**FR-W1:** A pattern is \"weak\" if `errors / attempts ≥ 0.5` AND `attempts ≥ 2` over rolling history (not single-test).")
bullet("**FR-W2:** Single-test results contribute to rolling history; rolling history is preserved across sessions in LocalStorage.")
bullet("**FR-W3:** A pattern graduates from \"weak\" to \"mastered\" after 4 consecutive correct answers across Drill or Test.")
bullet("**FR-W4:** Manual override: learner can mark a pattern \"known\" from Chapter 1, which sets it to mastered until the next miss.")

h2("6.7 NEW: SRS State Management")
bullet("**FR-S1:** Track per-pattern SRS box: `1d` / `3d` / `7d` / `14d` / `graduated`.")
bullet("**FR-S2:** Surface patterns in Drill where `nextDue ≤ now`.")
bullet("**FR-S3:** Do NOT preemptively drill patterns with no missed history.")
bullet("**FR-S4:** Wrong answer in Drill resets the pattern to box `1d` and `consecutiveCorrect = 0`.")

# ---------- 7) Data Model ----------
h1("7) Data Model (JSON-First)")

h2("7.1 Grammar Pattern Object (UPDATED)")
code_block('''{
  "id": "n5-001",
  "pattern": "〜です／〜ます",
  "category": "copula",
  "categoryOrder": 1,
  "patternOrder": 1,
  "meaning_en": "Polite copula / polite verb ending",
  "meaning_ja": "ていねいな いいかた",
  "form_rules": {
    "attaches_to": ["noun", "na_adjective", "verb_stem"],
    "conjugations": [
      { "form": "affirmative_present", "example": "がくせいです" },
      { "form": "negative_present",    "example": "がくせいじゃありません" },
      { "form": "affirmative_past",    "example": "がくせいでした" },
      { "form": "negative_past",       "example": "がくせいじゃありませんでした" }
    ]
  },
  "explanation_en": "Used in polite speech. です for nouns/adjectives; ます for verbs.",
  "examples": [
    { "ja": "わたしは がくせいです。",            "form": "affirmative", "furigana": [] },
    { "ja": "わたしは がくせいじゃありません。", "form": "negative",    "furigana": [] },
    { "ja": "あなたは がくせいですか。",          "form": "question",    "furigana": [] },
    { "ja": "きのうは やすみでした。",            "form": "past",        "furigana": [] }
  ],
  "common_mistakes": [
    {
      "wrong": "わたしは がくせいです じゃありません。",
      "right": "わたしは がくせいじゃありません。",
      "why":   "Don't combine です and じゃありません. Use one or the other for tense/polarity."
    }
  ],
  "contrasts": [
    { "with_pattern_id": "n5-XXX", "note": "..." }
  ],
  "notes": ""
}''')

h2("7.2 Question Object (UPDATED)")
p("MCQ shape:")
code_block('''{
  "id": "q-0001",
  "grammarPatternId": "n5-010",
  "type": "mcq",
  "direction": "j_to_e",
  "prompt_ja": "(  )に はいる いちばん いい ものを えらんでください。",
  "question_ja": "わたしは きのう ともだち(  )あいました。",
  "choices": ["を", "に", "で", "が"],
  "correctAnswer": "に",
  "explanation_en": "会う takes に to mark the person you meet.",
  "distractor_explanations": {
    "を": "を marks a direct object. 会う is intransitive - doesn't take を.",
    "で": "で marks location of action, not the person met.",
    "が": "が marks subject/new info, not the person met."
  },
  "high_confusion": true,
  "difficulty": 1
}''')

p("Sentence-order shape:")
code_block('''{
  "id": "q-0050",
  "grammarPatternId": "n5-005",
  "type": "sentence_order",
  "prompt_ja": "ただしい じゅんに ならべてください。",
  "tiles": ["わたしは", "がくせいです", "の", "だいがく"],
  "correctOrder": ["わたしは", "だいがく", "の", "がくせいです"],
  "translation_en": "I am a university student.",
  "explanation_en": "の links nouns; だいがくの がくせい = university's student."
}''')

h2("7.3 Test Result Object (UPDATED)")
code_block('''{
  "timestamp": "2026-04-29T19:30:00+09:00",
  "type": "test",
  "total": 20,
  "correct": 17,
  "incorrect": 3,
  "responses": [
    {
      "questionId": "q-0001",
      "userAnswer": "で",
      "isCorrect": false,
      "grammarPatternId": "n5-010"
    }
  ]
}''')

h2("7.4 NEW: Pattern History Object (LocalStorage)")
code_block('''{
  "n5-010": {
    "attempts": 7,
    "correct": 3,
    "incorrect": 4,
    "errorRate": 0.57,
    "isWeak": true,
    "isMastered": false,
    "isManuallyKnown": false,
    "lastSeen": "2026-04-29T19:30:00+09:00",
    "srsBox": "3d",
    "nextDue": "2026-05-02T19:30:00+09:00",
    "consecutiveCorrect": 0
  }
}''')

h2("7.5 NEW: Settings Object (LocalStorage)")
code_block('''{
  "furiganaOnN5Kanji": false,
  "lastTestLength": 20,
  "diagnosticCompleted": true,
  "lastDiagnosticDate": "2026-04-29T19:00:00+09:00"
}''')

# ---------- 8) System Architecture ----------
h1("8) System Architecture (LOCKED)")

h2("8.1 Tech Stack - Vanilla, No Build Step")
bullet("**HTML + CSS + Vanilla JavaScript** only.")
bullet("**No npm, no bundler, no toolchain.** Anyone with a modern browser can clone the repo, double-click `index.html`, and run.")
bullet("GitHub Pages serves the same files as the local file system.")
bullet("ES modules via `<script type=\"module\">` - modern browsers support imports natively.")
bullet("All paths relative.")
bullet("Hash routing for SPA refresh safety: `/#/learn`, `/#/test`, `/#/drill`, `/#/review`, `/#/summary`.")

h2("8.2 Why this stack")
bullet("Zero install for the learner - just a browser.")
bullet("Zero install for the content author - edit JSON / MD in any editor; refresh browser.")
bullet("Stable across Chrome / Edge / Firefox versions; no breaking-change surface from build tools.")
bullet("GitHub Pages compatible without any actions or workflows.")
bullet("No external CDN dependencies - everything is in the repo (resilient on flaky networks).")

h2("8.3 Content Pipeline (One-Time per Edit)")
bullet("Source-of-truth files: `KnowledgeBank/grammar_n5.md`, `KnowledgeBank/kanji_n5.md`, `KnowledgeBank/vocabulary_n5.md`, `KnowledgeBank/sources.md` (human-readable).")
bullet("App-consumed files: `/data/grammar.json`, `/data/questions.json`, `/data/n5_kanji_whitelist.json`, `/data/n5_vocab_whitelist.json`.")
bullet("`tools/build_data.py` - one-time conversion script (run by author after editing .md files). Generates JSON whitelists from kanji_n5.md and vocabulary_n5.md.")
bullet("`tools/check_coverage.py` - verifies every pattern in grammar_n5.md has an entry in grammar.json.")
bullet("`tools/lint_content.py` - flags non-N5 vocabulary or kanji in grammar.json / questions.json.")
bullet("These scripts are author-side only. The learner never runs Python.")

# ---------- 9) Repo Structure ----------
h1("9) Repository Structure")
code_block('''/JLPT/N5/                        # repo root for this app
  index.html                     # entry point - open this in browser
  /css/
    main.css
  /js/
    app.js                       # router + chapter coordinator
    learn.js                     # Chapter 1
    test.js                      # Chapter 2
    review.js                    # Chapter 3
    summary.js                   # Chapter 4
    drill.js                     # Drill / SRS engine
    diagnostic.js                # First-run diagnostic
    storage.js                   # LocalStorage adapter
    furigana.js                  # Furigana renderer + toggle
  /data/
    grammar.json                 # all N5 patterns (rich content)
    questions.json               # question bank
    n5_kanji_whitelist.json      # generated from kanji_n5.md
    n5_vocab_whitelist.json      # generated from vocabulary_n5.md
  /tools/
    build_data.py                # author tool - md to json whitelists
    check_coverage.py            # author tool - pattern coverage check
    lint_content.py              # author tool - N5 scope lint
  KnowledgeBank/grammar_n5.md                  # source-of-truth pattern catalog
  KnowledgeBank/kanji_n5.md                    # source-of-truth N5 kanji list
  KnowledgeBank/vocabulary_n5.md               # source-of-truth N5 vocab list
  KnowledgeBank/sources.md                     # reference / authority documentation
  JLPT N5 Grammar Tutor – Functional Spec.docx   # this file
  README.md''')

# ---------- 10) Deployment ----------
h1("10) Deployment Requirements (GitHub Pages)")
bullet("**DR-1:** Static site - no build step required.")
bullet("**DR-2:** Deploy via `/docs` folder OR `gh-pages` branch (either works for vanilla).")
bullet("**DR-3:** All paths relative.")
bullet("**DR-4:** Hash routing for refresh safety.")
bullet("**DR-5:** No external CDN dependencies - every asset self-hosted in repo.")
bullet("**DR-6:** Service worker (optional Phase 1.5) for full offline capability after first load.")

# ---------- 11) Non-Functional ----------
h1("11) Non-Functional Requirements")
bullet("**Performance:** First load under 2s on typical broadband; subsequent loads near-instant from cache.")
bullet("**Accessibility:** Full keyboard navigation; visible focus rings; WCAG AA contrast; skip-to-main-content link.")
bullet("**Compatibility:** Latest Chrome / Edge / Firefox / Safari on PC. Mobile-friendly is non-goal (Phase 2).")
bullet("**Reliability:** No data loss without explicit \"Reset progress\" confirmation.")
bullet("**Content integrity:** `tools/lint_content.py` flags non-N5 kanji or vocab before commit.")
bullet("**Offline:** After first load, app continues to work offline (everything served from filesystem / LocalStorage).")
bullet("**No telemetry / no tracking** - no analytics calls, no external resources.")

# ---------- 12) Acceptance ----------
h1("12) Validation & Acceptance Criteria (Must Pass)")

h2("Lesson Acceptance (Chapter 1)")
bullet("TOC lists every pattern in `KnowledgeBank/grammar_n5.md`, grouped by its 23 categories.")
bullet("Every pattern page shows the 7-block expanded template (§4.3).")
bullet("Furigana default is OFF on N5 kanji and ON for out-of-scope kanji; toggle persists across sessions.")
bullet("\"Mark as known\" checkbox works and updates the rolling history.")

h2("Test Acceptance (Chapter 2)")
bullet("Submit visible-but-disabled until all questions answered, with tooltip + remaining-count.")
bullet("Submit click immediately renders score / answer review / weak list.")
bullet("Sentence-ordering question type works (click-to-order tiles, array-equality grading).")
bullet("Question sampling balances across categories (max ceil(N/5) per category, N ≥ 8).")

h2("Diagnostic (NEW)")
bullet("First-run prompt is shown once and skippable.")
bullet("10 questions, sampled across the highest-frequency categories.")
bullet("Results seed the weak list but do not count toward score history.")

h2("Drill (NEW)")
bullet("Patterns due today shown with a badge count on the Drill tab.")
bullet("1d / 3d / 7d / 14d intervals advance correctly on right answers; reset to 1d on wrong.")
bullet("Pattern graduates after 4 consecutive correct.")
bullet("Drill session shows immediate feedback per question (not batched).")

h2("Review Acceptance (Chapter 3)")
bullet("Weak patterns shown by rolling threshold (≥ 50% error AND ≥ 2 attempts).")
bullet("Each weak pattern includes form rules + common mistakes + 4-6 fresh examples (different from Chapter 1).")
bullet("Per-distractor explanations shown for missed high-confusion questions.")

h2("Summary Acceptance (Chapter 4)")
bullet("Mastered / weak / untested all visible.")
bullet("Heatmap-style category grid renders.")
bullet("Suggested next step is actionable (links to the right chapter / mode).")

h2("Content Integrity")
bullet("`tools/check_coverage.py` exits 0 - every pattern in `KnowledgeBank/grammar_n5.md` has a matching `/data/grammar.json` entry.")
bullet("`tools/lint_content.py` exits 0 - no out-of-scope kanji/vocab in JSON without explicit furigana annotation.")

# ---------- 13) Future ----------
h1("13) Future Enhancements (Phase 2+)")
bullet("Adaptive testing with Anki-style ease factors (replaces Leitner-light).")
bullet("Audio on hover for furigana / pattern reading (TTS or recorded).")
bullet("Print / PDF export of lesson chapters.")
bullet("Multi-learner profiles (separate LocalStorage namespaces; or backend).")
bullet("N4 expansion (separate content set; same engine).")
bullet("Service worker for full offline + installable PWA.")
bullet("Mobile-friendly responsive layout.")
bullet("Listening practice (out of Phase 1 scope).")

# ---------- 14) Implementation Notes ----------
h1("14) Implementation Notes (Developer Guidance)")
bullet("Prefer dropdown fill-in-blank over free text to avoid ambiguous grading.")
bullet("Sentence-ordering tiles must be word/phrase chunks, NOT individual particles. Don't make the learner rearrange `わ・た・し・は` - that's grading kana, not grammar.")
bullet("Every question MUST be tagged with exactly one `grammarPatternId` to support gap analysis.")
bullet("Mark `high_confusion: true` only when ≥ 2 distractors are pedagogically defensible (worth explaining); otherwise leave `distractor_explanations` empty.")
bullet("LocalStorage key namespace: `jlpt-n5-tutor:*`.")
bullet("Furigana renderer: use `<ruby>...<rt>...</rt></ruby>` for true semantic ruby, not parentheses.")
bullet("All Japanese strings in code/data must be UTF-8 (no Shift-JIS / cp932).")
bullet("When generating Drill sessions, prefer questions the learner has missed before - but if none exist for a due pattern, sample any question for that pattern.")
bullet("Keep `js/` modules small (one chapter / mode per file). Aim < 300 LOC per file before splitting.")

# ============================================================
# v4: append v3.1 supplement (gap-fill addendum)
# ============================================================
# The supplement lives as markdown at specifications/JLPT-N5-Functional-
# Spec-v3.1-supplement.md. We render it inline so v4 is a single docx
# with all of v3's content PLUS the new sections (document control,
# glossary, RACI, user stories, KPIs, NFRs, test strategy, risks,
# open-questions, maintenance, release process, errata, audit protocol).

def _add_inline_runs(para, text, bold=False, italic=False):
    """Render bold (**...**), inline-code (`...`), and links [text](url)
    by toggling runs as we walk the string. Anything else is plain text."""
    i = 0
    while i < len(text):
        if text.startswith('**', i):
            # bold span
            end = text.find('**', i + 2)
            if end == -1:
                _add_inline_runs(para, text[i + 2:], bold=True, italic=italic); return
            inner = text[i + 2:end]
            r = para.add_run(inner); r.bold = True; r.italic = italic
            i = end + 2
        elif text.startswith('`', i):
            end = text.find('`', i + 1)
            if end == -1:
                r = para.add_run(text[i + 1:]); r.font.name = 'Consolas'; r.font.size = Pt(9); return
            r = para.add_run(text[i + 1:end]); r.font.name = 'Consolas'; r.font.size = Pt(9)
            i = end + 1
        elif text.startswith('[', i):
            close = text.find('](', i + 1)
            close2 = text.find(')', close + 2) if close != -1 else -1
            if close == -1 or close2 == -1:
                r = para.add_run(text[i]); r.bold = bold; r.italic = italic
                i += 1
                continue
            label = text[i + 1:close]
            r = para.add_run(label); r.bold = bold; r.italic = italic
            r.font.color.rgb = RGBColor(0x14, 0x45, 0x2A); r.underline = True
            i = close2 + 1
        else:
            # collect plain run until the next special marker
            stops = [text.find('**', i), text.find('`', i), text.find('[', i)]
            stops = [s for s in stops if s != -1]
            end = min(stops) if stops else len(text)
            chunk = text[i:end]
            if chunk:
                r = para.add_run(chunk); r.bold = bold; r.italic = italic
            i = end


def _emit_table(rows):
    """rows is a list of list-of-strings. First row is header."""
    if not rows: return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            # The default cell paragraph is empty - reuse it
            cell.paragraphs[0].text = ''
            _add_inline_runs(cell.paragraphs[0], cell_text,
                              bold=(r_idx == 0))
    doc.add_paragraph()  # spacer after table


def render_markdown(md_path):
    """Walk the supplement markdown line-by-line and emit doc nodes
    using the existing helpers (h1/h2/h3/p/bullet) plus an inline
    table renderer.

    Markdown features handled:
      # / ## / ### headings
      paragraphs (multi-line collapsed into one paragraph)
      - bullets (single level)
      | a | b | tables (with separator row skipped)
      **bold** / `code` / [text](url) inline
      --- horizontal rule
    """
    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # Skip blank lines (the helpers add their own spacing)
        if not stripped:
            i += 1
            continue
        # Headings
        if stripped.startswith('#### '):
            # Treat h4+ as bold paragraph (we don't use h4 in this doc)
            para = doc.add_paragraph()
            r = para.add_run(stripped[5:].strip()); r.bold = True
            i += 1; continue
        if stripped.startswith('### '):
            h3(stripped[4:].strip()); i += 1; continue
        if stripped.startswith('## '):
            h2(stripped[3:].strip()); i += 1; continue
        if stripped.startswith('# '):
            h1(stripped[2:].strip()); i += 1; continue
        # Horizontal rule
        if stripped in ('---', '***', '___'):
            hr(); i += 1; continue
        # Table: a line starting with | and the next line is a separator |---|---|
        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|?\s*:?-+', lines[i + 1].strip()):
            rows = []
            # header row
            header = [c.strip() for c in stripped.strip('|').split('|')]
            rows.append(header)
            i += 2  # skip header + separator
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            _emit_table(rows)
            continue
        # Bullet
        if stripped.startswith('- ') or stripped.startswith('* '):
            text_after = stripped[2:].strip()
            # Use the existing bullet() helper (it handles **bold**)
            # but we need inline-code + link too. Use add_paragraph + _add_inline_runs.
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            _add_inline_runs(para, text_after)
            i += 1; continue
        # Numbered list (1. text)
        if re.match(r'^\d+\.\s', stripped):
            text_after = re.sub(r'^\d+\.\s+', '', stripped)
            para = doc.add_paragraph(style='List Number')
            _add_inline_runs(para, text_after)
            i += 1; continue
        # Code fence
        if stripped.startswith('```'):
            # Collect until closing fence
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                block.append(lines[i]); i += 1
            i += 1  # consume closing fence
            code_block('\n'.join(block))
            continue
        # Regular paragraph: collect until blank line
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].strip().startswith(('#', '-', '*', '|', '>', '```'))
            or re.match(r'^\d+\.\s', lines[i].strip())
        ):
            para_lines.append(lines[i].strip())
            i += 1
        para = doc.add_paragraph()
        _add_inline_runs(para, ' '.join(para_lines))


import re  # used by render_markdown

SUPPLEMENT_PATH = ROOT / "specifications" / "JLPT-N5-Functional-Spec-v3.1-supplement.md"
if SUPPLEMENT_PATH.exists():
    hr()
    p("APPENDIX - v3.1 supplement merged into v4", bold=True, italic=True)
    p("(Source: specifications/JLPT-N5-Functional-Spec-v3.1-supplement.md)", italic=True)
    p("This appendix carries the gap-fill content that elevated v3 to v4: document control template, glossary, RACI, user stories, KPIs, full NFRs, test strategy, risks register, open-questions log, maintenance model, release process, errata against v3 sections, and the Pass-N audit protocol. Sections lettered A-E follow.")
    hr()
    render_markdown(SUPPLEMENT_PATH)

doc.save(OUT_PATH)
# stdout on Windows defaults to cp932, which can't encode the en-dash in OUT_PATH.
# Use ASCII-only for the success line so the script never trips on its own filename.
print("Saved spec docx ({} bytes).".format(Path(OUT_PATH).stat().st_size))
